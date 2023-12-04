import requests
from docx import Document
from docx.shared import Pt
from docx.opc.exceptions import PackageNotFoundError
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import openai
from openai import OpenAI
import time



#from googletrans import Translator
client = OpenAI(api_key="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")


def improve_grammar(text):
    try:
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": f"Please correct any grammatical or stylistic errors in the following text without altering the original message: {text}",
                }
            ],
            model="gpt-3.5-turbo",
        )
        text_content = chat_completion.choices[0].message.content
        return text_content
    except openai.OpenAIError as e:
        print(f"An OpenAI error occurred: {e}")


def translate_to_english(text):
    response = requests.post("http://localhost:5006/translate", json={"text": text, "src": "de", "dest": "en"})
    return response.json()["translated_text"]

def add_tab_to_sentences(text):
    return '\t' + text


try:
    # Initialize Document objects for reading and writing
    input_document = Document(r"C:\PersonalFiles - Copy\pythonProject\4Seiten.docx")
except PackageNotFoundError as e:
    print(f"PackageNotFoundError: {e}")
    exit()

# Initialize a new Document object for writing
output_document = Document()


try:
    # Loop through each paragraph in the input document, the translation is done by paragraph
    for para in input_document.paragraphs:

        # Skip footnotes based on font size and the presence of a line, if footnotes are sized 10, change accordingly
        if any(run.font.size == Pt(10) for run in para.runs):
            continue

        # Create a new paragraph in the output document
        try:
            translated_text = translate_to_english(para.text)
            if len(translated_text) > 50:  #for bug fixing and error hand
                print("Translated Text")
                print(translated_text)
                print("Fixed by spellcheck and style-check ")
                translated_text = improve_grammar(translated_text)
                print(translated_text)
                print("")

            new_text = add_tab_to_sentences(translated_text)
            new_para = output_document.add_paragraph()
            new_para.style = para.style
            new_para.alignment = para.alignment

            new_run = new_para.add_run(new_text)

            first_run = para.runs[0] if para.runs else None
            if first_run:
                new_run.bold = first_run.bold
                new_run.italic = first_run.italic
                new_run.underline = first_run.underline
                if first_run.font.size:
                    new_run.font.size = Pt(14) if first_run.font.size.pt in [11, 12] else first_run.font.size
                else:
                    new_run.font.size = Pt(14)
                new_run.font.name = 'Times New Roman'  # Set font to chosen font
        except Exception as e:
            print(f"An error occurred : {e}")
except Exception as e:
    print(f"An error occurred while processing the paragraphs: {e}")

try:
    # Save the output document
    output_document.save(r"C:\PersonalFiles - Copy\pythonProject\fertig4Seiten.docx")  # Specify the full path with a file name
except Exception as e:
    print(f"An error occurred while saving the document: {e}")
